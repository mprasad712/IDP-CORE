"""PaddleOCR wrapper for the IDP pipeline.

PDF pages are OCR'd with PaddleOCR — all pages are treated as scanned (no digital/scanned heuristic).
A page that is a single full-page image (a scanned page / a photo pasted into a PDF) is OCR'd from the
embedded image's OWN pixels; other pages are rendered to a raster first. This matters: re-rendering an
image-only page at a low DPI throws away resolution the embedded image already has, which is why an ID
card read perfectly as a raw image but became garbage once merged into a PDF (see _dominant_page_image).

Spreadsheet and Word files are extracted natively (no OCR needed).
"""

import multiprocessing
import platform
import sys
import cv2
import numpy as np
import fitz
import os
os.environ["PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK"] = "True"
from loguru import logger

# Maximum pixels on the long side sent to OCR.
# 1000px is the sweet spot for cheque/document text: 2.4x faster than 1500px
# with no meaningful accuracy loss for standard font sizes.
# Override with OCR_MAX_SIDE env var if a specific flow needs higher resolution.
_OCR_MAX_SIDE = int(os.environ.get("OCR_MAX_SIDE", "1000"))

# DPI used to rasterize a PDF PAGE when it must be rendered (text/mixed pages). Only used as a fallback:
# a page that is a single dominant image is OCR'd from the embedded image's OWN pixels instead (below).
_PDF_OCR_DPI = int(os.environ.get("OCR_PDF_DPI", "100"))

# A single embedded image covering at least this fraction of the page area IS the page (a scanned page /
# a photo pasted into a PDF). Such a page is OCR'd from the embedded image at its native resolution, not
# from a low-DPI page raster — rendering an A4 page at 100 DPI and capping to 1000px hands OCR a fraction
# of the pixels the original image had, which is exactly why a 720x1600 ID card read perfectly as a raw
# .jpeg but became garbage once merged into a PDF (same image, re-rendered smaller). See _dominant_page_image.
_FULL_PAGE_IMAGE_RATIO = float(os.environ.get("OCR_FULL_PAGE_IMAGE_RATIO", "0.6"))

# Only prefer the native embedded image when it actually carries resolution. Below this a low-res image
# stretched across a page benefits from the upscaled page render, so fall back to rendering.
_MIN_NATIVE_OCR_SIDE = int(os.environ.get("OCR_MIN_NATIVE_SIDE", "800"))

# Thread count: 4 is the sweet spot on most CPUs — beyond 4 the OCR pipeline
# has single-threaded bottlenecks and additional cores don't help.
_OCR_THREADS = int(os.environ.get("OCR_THREADS", "4"))

# MKL-DNN (OneDNN) gives ~5-10x speedup on Intel CPUs.
# On Windows + PaddlePaddle 3.x PIR it crashes with
# "ConvertPirAttribute2RuntimeAttribute not support ArrayAttribute<DoubleAttribute>".
# On Linux (Docker/AKS) it works correctly.
_MKLDNN_DEFAULT = sys.platform != "win32"

_paddle_ocr_available = False

try:
    from paddleocr import PaddleOCR
    _paddle_ocr_available = True
except ImportError:
    logger.warning("[OCR] paddleocr not found. Will use fallback.")
except Exception as e:
    logger.warning(f"[OCR] Failed to import paddleocr: {e}. Will use fallback.")

# Cache PaddleOCR instances per language
_ocr_instances: dict = {}


def get_ocr_instance(lang: str):
    global _paddle_ocr_available, _ocr_instances
    if not _paddle_ocr_available:
        return None

    ocr_lang = "hi" if lang.lower().strip() in ("hi", "hindi", "mixed") else "en"

    if ocr_lang not in _ocr_instances:
        enable_mkldnn = _MKLDNN_DEFAULT
        try:
            _ocr_instances[ocr_lang] = PaddleOCR(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
                lang=ocr_lang,
                enable_mkldnn=enable_mkldnn,
                cpu_threads=_OCR_THREADS,
            )
            mkldnn_str = "on" if enable_mkldnn else "off"
            logger.info(
                f"[OCR] Initialized PaddleOCR (lang={ocr_lang}, mkldnn={mkldnn_str}, "
                f"threads={_OCR_THREADS}, max_side={_OCR_MAX_SIDE})"
            )
        except Exception as e:
            logger.error(f"[OCR] Failed to initialize PaddleOCR ({ocr_lang}): {e}")
            return None

    return _ocr_instances[ocr_lang]


def _cap_image(img: np.ndarray) -> np.ndarray:
    """Downscale img so its longest side is at most _OCR_MAX_SIDE pixels."""
    h, w = img.shape[:2]
    if max(h, w) <= _OCR_MAX_SIDE:
        return img
    scale = _OCR_MAX_SIDE / max(h, w)
    return cv2.resize(img, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_AREA)


# ─────────────────────────────────────────────────────────────────────────────
# Office-format native extraction (no OCR)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_spreadsheet_text(file_bytes: bytes) -> list[dict]:
    import io
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    except Exception as e:
        logger.error(f"[OCR] Failed to read spreadsheet: {e}")
        return []
    results = []
    try:
        for sheet_index, ws in enumerate(wb.worksheets, start=1):
            for row in ws.iter_rows():
                for cell in row:
                    val = cell.value
                    if val is None or str(val).strip() == "":
                        continue
                    results.append({
                        "text": str(val).strip(),
                        "bounding_box": None,
                        "confidence": 1.0,
                        "page_number": sheet_index,
                    })
    finally:
        wb.close()
    return results


def _extract_docx_text(file_bytes: bytes) -> list[dict]:
    import io
    try:
        import docx
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"[OCR] Failed to read docx: {e}")
        return []
    results = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if text:
            results.append({"text": text, "bounding_box": None, "confidence": 1.0, "page_number": 1})
    for table in document.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                text = (cell.text or "").strip()
                if text:
                    results.append({"text": text, "bounding_box": None, "confidence": 1.0, "page_number": 1})
    return results


# ─────────────────────────────────────────────────────────────────────────────
# Core OCR helpers
# ─────────────────────────────────────────────────────────────────────────────

def _ocr_image(img: np.ndarray, ocr_model, page_number: int) -> list[dict]:
    """Run PaddleOCR on a numpy BGR image and return token dicts."""
    results = []
    try:
        ocr_res = ocr_model.ocr(img)
        if ocr_res and ocr_res[0]:
            for line in ocr_res[0]:
                box = line[0]       # [[x0,y0],[x1,y1],[x2,y2],[x3,y3]]
                text, conf = line[1]
                results.append({
                    "text": text,
                    "bounding_box": box,
                    "confidence": float(conf),
                    "page_number": page_number,
                })
    except Exception as e:
        logger.warning(f"[OCR] PaddleOCR failed on page {page_number}: {e}")
    return results


def _dominant_page_image(doc, page):
    """The page's single dominant embedded raster at its NATIVE resolution, or ``None``.

    A scanned page (an ID card, a photo, a faxed sheet) is one full-page image sitting on the page. The
    default path renders that page to a low-DPI pixmap and OCRs the pixmap — throwing away the embedded
    image's real resolution. A 720x1600 Aadhaar photo pasted into an A4 page, rendered at 100 DPI and
    capped to 1000px, reaches OCR at a fraction of its original detail; the SAME image OCR'd directly as a
    ``.jpeg`` reads perfectly. So when exactly one image covers most of the page, OCR its own bytes.

    Returns ``None`` — i.e. "render the page instead" — for:
      * pages with 0 images (vector/text pages) or >1 image (logo + content, a collage),
      * an image that does not dominate the page (a figure/logo, not the whole scan),
      * a low-resolution image that would gain nothing over the upscaled page render,
      * anything that fails to decode (exotic colorspace, JPEG2000) — never break OCR.
    """
    try:
        imgs = page.get_images(full=True)
        if len(imgs) != 1:
            return None
        xref = imgs[0][0]
        rects = page.get_image_rects(xref)
        if not rects:
            return None
        page_area = abs(page.rect.width * page.rect.height)
        img_area = sum(abs(r.width * r.height) for r in rects)
        if page_area <= 0 or (img_area / page_area) < _FULL_PAGE_IMAGE_RATIO:
            return None
        base = doc.extract_image(xref)
        arr = np.frombuffer(base["image"], np.uint8)
        img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if img is None or max(img.shape[:2]) < _MIN_NATIVE_OCR_SIDE:
            return None
        # The raw image is stored unrotated; honour any rotation the page applies to it so OCR sees it
        # upright. (The Scan Corrector usually bakes this in already, so rotation is normally 0.)
        rot = int(getattr(page, "rotation", 0) or 0) % 360
        if rot == 90:
            img = cv2.rotate(img, cv2.ROTATE_90_CLOCKWISE)
        elif rot == 180:
            img = cv2.rotate(img, cv2.ROTATE_180)
        elif rot == 270:
            img = cv2.rotate(img, cv2.ROTATE_90_COUNTERCLOCKWISE)
        return img
    except Exception as e:  # noqa: BLE001 — detection must never break the OCR path
        logger.debug(f"[OCR] dominant-image detection skipped: {e}")
        return None


def _pdf_to_images(file_bytes: bytes) -> list:
    """Convert PDF bytes to a list of PIL Images using pdf2image (poppler)."""
    try:
        from pdf2image import convert_from_bytes
        return convert_from_bytes(file_bytes)
    except Exception as e:
        logger.warning(f"[OCR] pdf2image failed ({e}), falling back to fitz pixmap")
        import fitz
        images = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            png_bytes = pix.tobytes("png")
            nparr = np.frombuffer(png_bytes, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            if img is not None:
                images.append(img)
        doc.close()
        return images


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

async def run_paddle_ocr(file_bytes: bytes, file_type: str, lang: str = "en") -> list[dict]:
    """OCR ``file_bytes`` and return a flat list of word tokens.

    Each token: ``{text, bounding_box, confidence, page_number}``

    - PDF: single-image pages OCR'd from the embedded image at native resolution; other pages rendered
      to a raster → PaddleOCR (no scanned/digital check)
    - Images: PaddleOCR directly
    - XLSX/DOCX: native text extraction, no OCR

    Runs the blocking OCR work in a thread executor so the event loop is never frozen.
    """
    import asyncio
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _paddle_ocr_impl, file_bytes, file_type, lang)


def _paddle_ocr_impl(file_bytes: bytes, file_type: str, lang: str = "en") -> list[dict]:
    """Synchronous OCR — called via run_in_executor; never call directly from async code."""
    file_type = file_type.lower().strip(".")
    is_pdf = (file_type == "pdf")

    if file_type in ("xlsx", "xls"):
        return _extract_spreadsheet_text(file_bytes)
    if file_type == "docx":
        return _extract_docx_text(file_bytes)

    ocr_model = get_ocr_instance(lang)

    if ocr_model is not None:
        try:
            results = []
            
            def parse_result(ocr_res_raw, page_num):
                parsed_lines = []
                # Check for new dict structure: [{'rec_texts': ..., 'dt_polys': ..., 'rec_scores': ...}]
                if isinstance(ocr_res_raw, list) and len(ocr_res_raw) > 0 and isinstance(ocr_res_raw[0], dict) and "rec_texts" in ocr_res_raw[0]:
                    data = ocr_res_raw[0]
                    texts = data.get("rec_texts", [])
                    boxes = data.get("dt_polys", []) or data.get("rec_boxes", [])
                    scores = data.get("rec_scores", [])
                    for idx in range(min(len(texts), len(boxes), len(scores))):
                        box = boxes[idx]
                        if isinstance(box, np.ndarray):
                            box = box.tolist()
                        parsed_lines.append({
                            "text": str(texts[idx]),
                            "bounding_box": box,
                            "confidence": float(scores[idx]),
                            "page_number": page_num
                        })
                else:
                    # Legacy structure: list of list of elements like [box, (text, conf)]
                    if ocr_res_raw and isinstance(ocr_res_raw, list) and isinstance(ocr_res_raw[0], list):
                        for line in ocr_res_raw[0]:
                            box = line[0]
                            if isinstance(box, np.ndarray):
                                box = box.tolist()
                            text, conf = line[1]
                            parsed_lines.append({
                                "text": str(text),
                                "bounding_box": box,
                                "confidence": float(conf),
                                "page_number": page_num
                            })
                return parsed_lines

            if is_pdf:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page_num in range(len(doc)):
                    page = doc[page_num]

                    # Scanned page = one full-page image: OCR its OWN pixels, not a low-DPI page raster
                    # (see _dominant_page_image — this is what makes a merged ID card read as well as the
                    # standalone image). Everything else renders the page as before.
                    native = _dominant_page_image(doc, page)
                    if native is not None:
                        img = native
                        src = "embedded image @native"
                    else:
                        pix = page.get_pixmap(dpi=_PDF_OCR_DPI)
                        page_img_bytes = pix.tobytes("png")
                        nparr = np.frombuffer(page_img_bytes, np.uint8)
                        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                        src = f"page render @{_PDF_OCR_DPI}dpi"

                    if img is None:
                        continue

                    _pre = img.shape[:2]
                    img = _cap_image(img)
                    h, w = img.shape[:2]
                    logger.info(f"[OCR] running page {page_num + 1} via {src} ({_pre[1]}×{_pre[0]}→{w}×{h}px)")
                    try:
                        ocr_res = list(ocr_model.predict(img))
                    except Exception:
                        ocr_res = ocr_model.ocr(img)

                    page_tokens = parse_result(ocr_res, page_num + 1)
                    logger.info(f"[OCR] page {page_num + 1} done — {len(page_tokens)} tokens")
                    results.extend(page_tokens)
                doc.close()
            else:
                nparr = np.frombuffer(file_bytes, np.uint8)
                img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                if img is not None:
                    img = _cap_image(img)
                    h, w = img.shape[:2]
                    logger.info(f"[OCR] running image ({w}×{h}px)")
                    try:
                        ocr_res = list(ocr_model.predict(img))
                    except Exception:
                        ocr_res = ocr_model.ocr(img)
                    page_tokens = parse_result(ocr_res, 1)
                    logger.info(f"[OCR] image done — {len(page_tokens)} tokens")
                    results.extend(page_tokens)
            return results
        except Exception as e:
            logger.warning(f"[OCR] PaddleOCR execution failed: {e}. Falling back to text/mock extraction.")

    # Fallback to PyMuPDF word extraction for PDFs, and basic mock for images
    results = []
    if is_pdf:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page_num in range(len(doc)):
                for w in doc[page_num].get_text("words"):
                    x0, y0, x1, y1, word_text, *_ = w
                    results.append({
                        "text": word_text,
                        "bounding_box": [[x0, y0], [x1, y0], [x1, y1], [x0, y1]],
                        "confidence": 1.0,
                        "page_number": page_num + 1,
                    })
            doc.close()
        except Exception as e:
            logger.error(f"[OCR] PDF native fallback failed: {e}")
        return results

    # ── Image ────────────────────────────────────────────────────────────────
    if ocr_model is not None:
        try:
            nparr = np.frombuffer(file_bytes, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            if img is not None:
                return _ocr_image(img, ocr_model, page_number=1)
        except Exception as e:
            logger.warning(f"[OCR] Image PaddleOCR failed: {e}")

    # Mock fallback for images when PaddleOCR unavailable
    try:
        nparr = np.frombuffer(file_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        h, w = (img.shape[:2] if img is not None else (100, 100))
        return [{
            "text": "Scanned Image (Fallback)",
            "bounding_box": [[0.0, 0.0], [float(w), 0.0], [float(w), float(h)], [0.0, float(h)]],
            "confidence": 0.8,
            "page_number": 1,
        }]
    except Exception as e:
        logger.error(f"[OCR] Image fallback failed: {e}")
        return []
