"""Digital-document text path for the IDP pipeline (PyMuPDF + office libs).

A digital PDF already carries a text layer, so it skips OCR entirely. This module
detects digital-vs-scanned and extracts native text + word tokens. Scanned documents
go to ``services/idp/ocr.py`` (PaddleOCR) instead. Kept independent of the OCR engine
because the native-text path has no need to load PaddleOCR.
"""

from __future__ import annotations

import io

from loguru import logger

DIGITAL = "digital"
SCANNED = "scanned"

_IMG_EXT = {"png", "jpg", "jpeg", "tiff", "tif", "bmp", "webp"}
_OFFICE_EXT = {"xlsx", "xls", "docx"}


def _norm_type(file_type: str) -> str:
    return (file_type or "").lower().lstrip(".").strip()


def is_pdf_searchable(
    file_bytes: bytes, file_type: str, min_words: int = 30, min_text_length: int = 50
) -> dict[str, str]:
    """Per-page digital/scanned classification.

    Returns ``{page_no(str): 'digital'|'scanned'}``. Images -> one scanned page;
    office formats -> one digital page (native text). A PDF page is ``digital`` when it
    has at least ``min_words`` extractable words OR at least ``min_text_length`` chars of
    extractable text (``min_text_length`` is the detector node's configurable knob).
    """
    ft = _norm_type(file_type)
    if ft in _IMG_EXT:
        return {"1": SCANNED}
    if ft in _OFFICE_EXT:
        return {"1": DIGITAL}
    if ft == "txt":
        return {"1": DIGITAL}
    if ft != "pdf":
        return {"1": SCANNED}

    import fitz

    status: dict[str, str] = {}
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            for i in range(len(doc)):
                page = doc[i]
                n_words = len(page.get_text("words"))
                n_chars = len((page.get_text("text") or "").strip())
                is_digital = n_words >= min_words or n_chars >= min_text_length
                status[str(i + 1)] = DIGITAL if is_digital else SCANNED
        finally:
            doc.close()
    except Exception as e:
        logger.warning(f"[text_layer] is_pdf_searchable failed: {e}")
        return {"1": SCANNED}
    return status or {"1": SCANNED}


def classify_document(
    file_bytes: bytes, file_type: str, min_text_length: int = 50, min_words: int = 30
) -> tuple[str, dict[str, str]]:
    """Return ``('digital'|'scanned'|'mixed', page_status)``."""
    page_status = is_pdf_searchable(
        file_bytes, file_type, min_words=min_words, min_text_length=min_text_length
    )
    labels = set(page_status.values())
    if labels == {DIGITAL}:
        overall = DIGITAL
    elif labels == {SCANNED}:
        overall = SCANNED
    else:
        overall = "mixed"
    return overall, page_status


def extract_native_text(file_bytes: bytes, file_type: str) -> tuple[str, list[dict]]:
    """Digital path: return ``(text, tokens)``.

    ``text``: pages joined by form-feed (``\\x0c``) so page-slicing works downstream.
    ``tokens``: ``[{text, bounding_box, confidence, page_number}]`` for source-location.
    """
    ft = _norm_type(file_type)
    if ft in ("xlsx", "xls"):
        return _office_xlsx(file_bytes)
    if ft == "docx":
        return _office_docx(file_bytes)
    if ft == "txt":
        return _plain_text(file_bytes)
    if ft != "pdf":
        return "", []

    import fitz

    pages_text: list[str] = []
    tokens: list[dict] = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            for i in range(len(doc)):
                page = doc[i]
                pages_text.append(page.get_text("text") or "")
                for w in page.get_text("words"):
                    x0, y0, x1, y1, word_text = w[0], w[1], w[2], w[3], w[4]
                    if not str(word_text).strip():
                        continue
                    tokens.append(
                        {
                            "text": word_text,
                            "bounding_box": [[x0, y0], [x1, y0], [x1, y1], [x0, y1]],
                            "confidence": 1.0,
                            "page_number": i + 1,
                        }
                    )
        finally:
            doc.close()
    except Exception as e:
        logger.error(f"[text_layer] extract_native_text failed: {e}")
        return "", []
    return "\x0c".join(pages_text), tokens


def _plain_text(file_bytes: bytes) -> tuple[str, list[dict]]:
    """Plain .txt: decode and emit one token per non-empty line. Form-feeds (\\x0c) split pages."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="replace")
    tokens: list[dict] = []
    for pidx, page in enumerate(text.split("\x0c"), start=1):
        for line in page.splitlines():
            if line.strip():
                tokens.append({"text": line, "bounding_box": None, "confidence": 1.0, "page_number": pidx})
    return text, tokens


def _office_xlsx(file_bytes: bytes) -> tuple[str, list[dict]]:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    except Exception as e:
        logger.error(f"[text_layer] xlsx read failed: {e}")
        return "", []

    tokens: list[dict] = []
    parts: list[str] = []
    try:
        for sidx, ws in enumerate(wb.worksheets, start=1):
            for row in ws.iter_rows():
                for cell in row:
                    val = cell.value
                    if val is None or str(val).strip() == "":
                        continue
                    s = str(val).strip()
                    parts.append(s)
                    tokens.append({"text": s, "bounding_box": None, "confidence": 1.0, "page_number": sidx})
    finally:
        wb.close()
    return "\n".join(parts), tokens


def _office_docx(file_bytes: bytes) -> tuple[str, list[dict]]:
    try:
        import docx

        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"[text_layer] docx read failed: {e}")
        return "", []

    tokens: list[dict] = []
    parts: list[str] = []
    for para in document.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
            tokens.append({"text": t, "bounding_box": None, "confidence": 1.0, "page_number": 1})
    for table in document.tables:
        for trow in table.rows:
            for cell in trow.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
                    tokens.append({"text": t, "bounding_box": None, "confidence": 1.0, "page_number": 1})
    return "\n".join(parts), tokens
